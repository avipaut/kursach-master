from flask import Blueprint, render_template, request, redirect, url_for, send_from_directory, flash, current_app,jsonify
import os
import unicodedata
import chardet
from flask_login import current_user, login_required
from werkzeug.utils import secure_filename
from datetime import datetime
import mimetypes
from deep_translator import GoogleTranslator
import re
from .config import BASE_UPLOAD_FOLDER 
documents_bp = Blueprint('documents', __name__)
from .trash import TrashManager  # добавьте этот импорт
from routes.models import User  # Добавим импорт модели User
from flask import send_file
import tempfile
from pdf2image import convert_from_bytes
from docx2pdf import convert
import io



from routes.auth import role_required  # Импортируем декоратор для проверки роли
# Основная папка для загрузок
trash_manager = TrashManager(BASE_UPLOAD_FOLDER)
import shutil
@documents_bp.route('/admin/users')
@role_required('admin')  # Только для администраторов
def list_users():
    users = User.query.all()
    return render_template('documents/admin_users.html', users=users)
def normalize_filename(filename):
    """
    Normalize Unicode filename while preserving Cyrillic characters
    """
    # Split filename and extension
    name, ext = os.path.splitext(filename)
    
    # Normalize the name part while preserving Cyrillic
    normalized = unicodedata.normalize('NFKC', name)
    
    # Secure the filename while keeping Cyrillic characters
    secured = secure_filename_with_cyrillic(normalized)
    
    # Return the normalized name with the original extension
    return f"{secured}{ext.lower()}"

def get_user_upload_folder(user_id):
    """Create and return a unique folder for each user's uploads"""
    user_folder = os.path.join(BASE_UPLOAD_FOLDER, str(user_id))
    os.makedirs(user_folder, exist_ok=True)
    return user_folder

def get_common_folder():
    """Create and return a folder for common documents"""
    common_folder = os.path.join(BASE_UPLOAD_FOLDER, "common")
    os.makedirs(common_folder, exist_ok=True)
    return common_folder
# Add these imports to your documents.py file
from docx import Document
import html
import re
import os
import tempfile
from bs4 import BeautifulSoup
import mammoth


@documents_bp.route('/send_to_chat', methods=['POST'])
@login_required
def send_to_chat():
    try:
        data = request.json
        filename = data.get('filename')
        category = data.get('category')
        recipients = data.get('recipients', [])
        message_text = data.get('message', '')

        if not filename or not category or not recipients:
            return jsonify({'success': False, 'error': 'Missing required parameters'}), 400

        # Получаем путь к файлу
        if category == 'personal':
            folder = get_user_upload_folder(current_user.id)
        elif category == 'common':
            folder = get_common_folder()
        elif category == 'important':
            folder = get_important_folder()
        else:
            return jsonify({'success': False, 'error': 'Invalid category'}), 400

        filepath = os.path.join(folder, filename)
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'error': 'File not found'}), 404

        # Создаем лобби с выбранными пользователями
        user_ids = [int(uid) for uid in recipients]
        user_ids.append(current_user.id)  # Добавляем текущего пользователя

        # Проверяем, есть ли уже существующее лобби между этими пользователями
        existing_lobby = Lobby.query.filter(
            Lobby.is_group.is_(False),
            Lobby.users.any(User.id.in_(user_ids))
        ).group_by(Lobby.id).having(
            db.func.count(User.id) == len(user_ids)
        ).first()

        if existing_lobby:
            lobby = existing_lobby
        else:
            # Создаем новое лобби
            lobby = Lobby(is_group=len(user_ids) > 2)
            for user_id in user_ids:
                user = User.query.get(user_id)
                if user:
                    lobby.users.append(user)
            db.session.add(lobby)
            db.session.commit()

        # Отправляем файл в чат
        file_size = os.path.getsize(filepath)
        file_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
        unique_filename = f"{uuid.uuid4().hex}_{filename}"
        destination_path = os.path.join(UPLOAD_FOLDER, unique_filename)

        # Копируем файл в папку загрузок чата
        shutil.copy(filepath, destination_path)

        # Создаем сообщение
        new_message = Message(
            sender_id=current_user.id,
            lobby_id=lobby.id,
            text=message_text,
            message_type=get_file_type(filename),
            file_path=f'/uploads/{unique_filename}',
            file_name=filename,
            file_size=file_size,
            file_type=file_type
        )
        db.session.add(new_message)
        db.session.commit()

        # Отправляем уведомления получателям
        for user_id in user_ids:
            if user_id != current_user.id:
                socketio.emit('new_message', new_message.to_dict(), room=f'user_{user_id}')

        return jsonify({'success': True, 'lobby_id': lobby.id})

    except Exception as e:
        current_app.logger.error(f"Error sending file to chat: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@documents_bp.route('/preview/<category>/<path:filename>')
@login_required
def preview_file(category, filename):
    from urllib.parse import unquote
    from flask import abort, current_app, send_from_directory, Response, render_template_string
    import os
    
    try:
        # Decode the filename from URL
        filename = unquote(filename)
        
        # Validate category
        if category not in ['personal', 'common', 'important']:
            abort(404, description="Invalid category")
            
        # Get folder path
        if category == 'personal':
            folder = get_user_upload_folder(current_user.id)
        elif category == 'common':
            folder = get_common_folder()
        else:
            folder = get_important_folder()
            
        # Full path to file
        filepath = os.path.join(folder, filename)
        
        # Check if file exists
        if not os.path.exists(filepath):
            current_app.logger.error(f"File not found: {filepath}")
            abort(404, description="File not found")
        
        # Get file extension
        file_ext = os.path.splitext(filename)[1].lower()
        
        # For PDF files use browser's built-in viewer
        if file_ext == '.pdf':
            response = send_from_directory(folder, filename)
            response.headers['X-Frame-Options'] = 'SAMEORIGIN'
            response.headers['Content-Security-Policy'] = "frame-ancestors 'self'"
            return response
        
        # For DOCX files, convert to HTML for preview
        elif file_ext in ['.docx', '.doc']:
            try:
                # Convert DOCX to HTML using mammoth
                with open(filepath, 'rb') as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html_content = result.value
                
                # Create a complete HTML document
                preview_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>{html.escape(filename)}</title>
                    <style>
                        body {{ 
                            font-family: Arial, sans-serif; 
                            margin: 40px; 
                            line-height: 1.6;
                        }}
                        img {{ max-width: 100%; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                
                return Response(preview_html, mimetype='text/html')
            except Exception as e:
                current_app.logger.error(f"Error converting DOCX to HTML: {str(e)}")
                # Fall back to download if conversion fails
                return send_from_directory(folder, filename, as_attachment=True)
        
        # For TXT files, display content
        elif file_ext == '.txt':
            try:
                encoding = detect_file_encoding(filepath)
                with open(filepath, 'r', encoding=encoding, errors='replace') as f:
                    content = f.read()
                
                # Create a simple HTML display
                preview_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>{html.escape(filename)}</title>
                    <style>
                        body {{ 
                            font-family: monospace; 
                            margin: 40px; 
                            white-space: pre-wrap;
                        }}
                    </style>
                </head>
                <body>
                    {html.escape(content)}
                </body>
                </html>
                """
                
                return Response(preview_html, mimetype='text/html')
            except Exception as e:
                current_app.logger.error(f"Error reading TXT file: {str(e)}")
                return send_from_directory(folder, filename, as_attachment=True)
        
        # For other file types
        else:
            response = send_from_directory(folder, filename)
            response.headers['X-Frame-Options'] = 'SAMEORIGIN'
            response.headers['Content-Security-Policy'] = "frame-ancestors 'self'"
            return response
        
    except Exception as e:
        current_app.logger.error(f"Error in preview_file: {str(e)}")
        abort(500, description="Internal server error")

def get_common_folder():
    """Create and return a folder for common documents with proper permissions"""
    common_folder = os.path.join(BASE_UPLOAD_FOLDER, "common")
    try:
        os.makedirs(common_folder, exist_ok=True)
        # Устанавливаем правильные права доступа
        os.chmod(common_folder, 0o755)
    except Exception as e:
        current_app.logger.error(f"Error creating common folder: {str(e)}")
        raise
    return common_folder

# В том же documents.py добавим вспомогательную функцию
def check_file_exists(category, filename):
    """Проверяет существование файла и возвращает его путь или None"""
    try:
        filename = unquote(filename)
        
        if category == 'personal':
            folder = get_user_upload_folder(current_user.id)
        elif category == 'common':
            folder = get_common_folder()
        elif category == 'important':
            folder = get_important_folder()
        else:
            return None
            
        filepath = os.path.join(folder, filename)
        return filepath if os.path.exists(filepath) else None
    except:
        return None

def get_important_folder():
    """Create and return a folder for important documents"""
    important_folder = os.path.join(BASE_UPLOAD_FOLDER, "important")
    os.makedirs(important_folder, exist_ok=True)
    return important_folder

translated = GoogleTranslator(source='auto', target='en').translate("Привет")
def allowed_file(filename):
    """Check if the file extension is allowed"""
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'doc', 'rtf', 'odt'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_icon(filename):
    """Determine the appropriate icon based on file type"""
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    icons = {
        'pdf': 'fa-file-pdf',
        'docx': 'fa-file-word',
        'doc': 'fa-file-word',
        'txt': 'fa-file-alt',
        'rtf': 'fa-file-alt',
        'odt': 'fa-file-alt'
    }
    return icons.get(ext, 'fa-file')

def detect_file_encoding(filepath):
    """Detect file encoding using chardet"""
    with open(filepath, 'rb') as file:
        raw_data = file.read()
        result = chardet.detect(raw_data)
    return result['encoding'] or 'utf-8'

def translate_file(filepath, direction='ky-ru'):
    """Улучшенный перевод файлов с обработкой ошибок"""
    try:
        # Проверка существования файла
        if not os.path.exists(filepath):
            return None, "Файл не существует"
        
        # Проверка размера файла
        if os.path.getsize(filepath) == 0:
            return None, "Файл пуст"
        
        # Извлечение текста
        try:
            content = extract_text_from_file(filepath)
            if not content or not content.strip():
                return None, "Не удалось извлечь текст или файл пуст"
        except Exception as e:
            return None, f"Ошибка извлечения текста: {str(e)}"
        
        # Настройки перевода
        lang_map = {
            'ky-ru': ('ky', 'ru'),
            'ru-ky': ('ru', 'ky')
        }
        
        if direction not in lang_map:
            return None, "Неверное направление перевода"
        
        source, target = lang_map[direction]
        
        # Разделение на части для больших файлов
        max_chunk_size = 3000  # Уменьшенный размер для надежности
        chunks = split_text_for_translation(content, max_chunk_size)
        translated = []
        
        for chunk in chunks:
            try:
                if chunk.strip():
                    # Добавляем задержку между запросами
                    import time
                    time.sleep(1)  # 1 секунда между запросами
                    
                    translator = GoogleTranslator(source=source, target=target)
                    translated_chunk = translator.translate(chunk)
                    translated.append(translated_chunk)
            except Exception as e:
                # Логирование ошибки
                print(f"Translation error for chunk: {str(e)}")
                return None, f"Ошибка перевода части текста: {str(e)}"
        
        return ' '.join(translated), None
    
    except Exception as e:
        return None, f"Ошибка обработки файла: {str(e)}"

from tenacity import retry, stop_after_attempt, wait_exponential

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def translate_with_retry(translator, text):
    return translator.translate(text)
def extract_text_from_file(filepath):
    """Универсальное извлечение текста из разных форматов файлов"""
    ext = os.path.splitext(filepath)[1].lower()
    
    try:
        # TXT файлы
        if ext == '.txt':
            with open(filepath, 'rb') as f:
                encoding = chardet.detect(f.read())['encoding']
            with open(filepath, 'r', encoding=encoding or 'utf-8', errors='replace') as f:
                return f.read()
        
        # PDF файлы
        elif ext == '.pdf':
            try:
                from pdfminer.high_level import extract_text
                return extract_text(filepath)
            except ImportError:
                from PyPDF2 import PdfReader
                return '\n'.join([page.extract_text() or '' for page in PdfReader(filepath).pages])
        
        # DOCX файлы
        elif ext == '.docx':
            try:
                from docx import Document
                return '\n'.join([p.text for p in Document(filepath).paragraphs if p.text])
            except ImportError:
                import zipfile
                from xml.etree.ElementTree import XML
                with zipfile.ZipFile(filepath) as z:
                    with z.open('word/document.xml') as f:
                        return ' '.join([n.text for n in XML(f.read()).iter() if n.text])
        
        # DOC файлы (требует win32com или antiword)
        elif ext == '.doc':
            try:
                from win32com.client import Dispatch
                word = Dispatch('Word.Application')
                doc = word.Documents.Open(os.path.abspath(filepath))
                text = doc.Content.Text
                word.Quit()
                return text
            except:
                import subprocess
                return subprocess.check_output(['antiword', filepath]).decode('utf-8', 'ignore')
        
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    except Exception as e:
        raise Exception(f"Failed to extract text: {str(e)}")
def split_text_for_translation(text, max_length=5000):
    """Разделение текста на части для перевода"""
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para_length = len(para)
        if current_length + para_length > max_length and current_chunk:
            chunks.append('\n'.join(current_chunk))
            current_chunk = []
            current_length = 0
            
        current_chunk.append(para)
        current_length += para_length
    
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks



@documents_bp.route('/translate_upload', methods=['GET', 'POST'])
@login_required
def translate_upload():
    if request.method == 'GET':
        return render_template('documents/translate_upload.html')
    
    # Обработка POST запроса
    if 'file' not in request.files:
        flash('No file selected', 'error')
        return redirect(request.url)

    file = request.files['file']
    translation_direction = request.form.get('direction', 'ky-ru')
    doc_category = request.form.get('category', 'personal')

    if not file or file.filename == '':
        flash('No file selected', 'error')
        return redirect(request.url)
        
    if not allowed_file(file.filename):
        flash('Unsupported file type', 'error')
        return redirect(request.url)

    # Нормализация имени файла
    filename = normalize_filename(file.filename)
    
    # Выбор папки для сохранения
    if doc_category == 'common':
        folder = get_common_folder()
    elif doc_category == 'important':
        folder = get_important_folder()
    else:
        folder = get_user_upload_folder(current_user.id)
        
    # Сохраняем оригинальный файл
    original_path = os.path.join(folder, filename)
    file.save(original_path)
    
    # Выполняем перевод
    translated_content, error = translate_file(original_path, direction=translation_direction)
    
    if error:
        flash(f'Translation failed: {error}', 'error')
        return redirect(url_for('documents.view_file', filename=filename, category=doc_category))
    
    # Сохраняем переведенный файл
    name, ext = os.path.splitext(filename)
    translated_filename = f"{name}_translated_{translation_direction}{ext}"
    translated_path = os.path.join(folder, translated_filename)
    
    try:
        if ext == '.txt':
            with open(translated_path, 'w', encoding='utf-8') as f:
                f.write(translated_content)
        elif ext == '.docx':
            save_as_docx(translated_content, translated_path)
        elif ext == '.pdf':
            save_as_pdf(translated_content, translated_path)
            
        flash('File translated successfully!', 'success')
        return redirect(url_for('documents.view_file', 
                             filename=translated_filename, 
                             category=doc_category))
    except Exception as e:
        flash(f'Error saving translated file: {str(e)}', 'error')
        return redirect(url_for('documents.view_file', filename=filename, category=doc_category))

@documents_bp.route('/')
@login_required
def documents():
    return redirect(url_for('documents.list_documents', category='personal'))

@documents_bp.route('/<category>')
@login_required
def list_documents(category):
    if category not in ['personal', 'common', 'important']:
        category = 'personal'
    
    if category == 'personal':
        folder = get_user_upload_folder(current_user.id)
    elif category == 'common':
        folder = get_common_folder()
    else:
        folder = get_important_folder()
    
    documents = []
    
    try:
        for filename in os.listdir(folder):
            filepath = os.path.join(folder, filename)
            if os.path.isfile(filepath):
                file_stats = os.stat(filepath)
                documents.append({
                    'name': filename,
                    'size': f"{file_stats.st_size / 1024:.2f} KB",
                    'uploaded_at': datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M'),
                    'icon': get_file_icon(filename),
                    'exists': True
                })
    except Exception as e:
        flash(f'Error reading documents: {str(e)}', 'error')

    if not documents:
        documents = [{'name': "No documents available. Upload a new file!", 'exists': False}]
    
    # Получаем всех пользователей для выпадающего списка
    from .models import User  # Добавляем импорт
    all_users = User.query.filter(User.id != current_user.id).all()
    
    return render_template('documents/documents.html', 
                         documents=documents, 
                         active_category=category,
                         check_file_exists=check_file_exists,
                         all_users=all_users)  # Передаем список пользователей

@documents_bp.route('/create_documents', methods=['GET', 'POST'])
def create_documents():
    # ваш код для создания документов
    return render_template('documents/create_document.html')

@documents_bp.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        flash('No file part', 'error')
        return redirect(request.url)
    
    file = request.files['file']
    if file.filename == '':
        flash('No selected file', 'error')
        return redirect(request.url)
    
    doc_category = request.form.get('category', 'personal')
    
    if file and allowed_file(file.filename):
        # Use our custom filename sanitizer instead of secure_filename
        filename = normalize_filename(file.filename)
        
        # Особая обработка для важных документов от обычных пользователей
        if doc_category == 'important' and not current_user.is_admin:
            # Создаем временную папку для хранения файлов до одобрения
            temp_folder = os.path.join(BASE_UPLOAD_FOLDER, 'pending')
            os.makedirs(temp_folder, exist_ok=True)
            
            # Генерируем уникальное имя файла для временного хранения
            unique_filename = f"{current_user.id}_{int(datetime.now().timestamp())}_{filename}"
            temp_filepath = os.path.join(temp_folder, unique_filename)
            
            # Сохраняем файл во временную папку
            file.save(temp_filepath)
            
            # Создаем запись о файле, ожидающем одобрения
            from .models import db, PendingFile
            pending_file = PendingFile(
                filename=unique_filename,
                original_filename=filename,
                user_id=current_user.id,
                temp_path=temp_filepath
            )
            db.session.add(pending_file)
            db.session.commit()
            
            # Отправляем уведомления всем администраторам
            from .models import User
            from .notifications import add_notification
            
            admins = User.query.filter_by(is_admin=True).all()
            for admin in admins:
                add_notification(
                    user_id=admin.id,
                    message=f'Пользователь {current_user.username} запрашивает загрузку файла "{filename}" в папку важных документов',
                    category='warning',
                    link=url_for('documents.review_pending_file', file_id=pending_file.id)
                )
            
            flash(f'Файл "{filename}" отправлен на одобрение администратору', 'info')
            return redirect(url_for('documents.list_documents', category='personal'))
        
        # Стандартная обработка для других категорий или для администраторов
        if doc_category == 'common':
            folder = get_common_folder()
        elif doc_category == 'important':
            folder = get_important_folder()
        else:
            folder = get_user_upload_folder(current_user.id)
            
        filepath = os.path.join(folder, filename)
        file.save(filepath)
        
        from .notifications import add_notification
        add_notification(
            user_id=current_user.id,
            message=f'Ваш файл "{filename}" успешно загружен',
            category='success',
            link=url_for('documents.view_file', filename=filename, category=doc_category)
        )
        
        flash(f'Файл {filename} загружен успешно!', 'success')
        return redirect(url_for('documents.list_documents', category=doc_category))
    
    flash('Недопустимый тип файла', 'error')
    return redirect(url_for('documents.list_documents', category='personal'))


@documents_bp.route('/admin/pending_files')
@login_required
@role_required('admin')
def pending_files():
    from .models import PendingFile
    
    pending_files = PendingFile.query.filter_by(status='pending').all()
    return render_template('documents/admin_pending_files.html', pending_files=pending_files)

@documents_bp.route('/admin/review_file/<int:file_id>')
@login_required
@role_required('admin')
def review_pending_file(file_id):
    from .models import PendingFile, User
    
    pending_file = PendingFile.query.get_or_404(file_id)
    user = User.query.get(pending_file.user_id)
    
    return render_template('documents/review_file.html', file=pending_file, user=user)

@documents_bp.route('/admin/approve_file/<int:file_id>', methods=['POST'])
@login_required
@role_required('admin')
def approve_file(file_id):
    from .models import db, PendingFile, User
    from .notifications import add_notification
    
    pending_file = PendingFile.query.get_or_404(file_id)
    
    if pending_file.status != 'pending':
        flash('Этот файл уже обработан', 'warning')
        return redirect(url_for('documents.pending_files'))
    
    try:
        # Получаем папку для важных документов
        important_folder = get_important_folder()
        
        # Перемещаем файл из временной папки в папку важных документов
        target_path = os.path.join(important_folder, pending_file.original_filename)
        shutil.copy(pending_file.temp_path, target_path)
        
        # Обновляем статус файла
        pending_file.status = 'approved'
        db.session.commit()
        
        # Отправляем уведомление пользователю
        add_notification(
            user_id=pending_file.user_id,
            message=f'Ваш файл "{pending_file.original_filename}" был одобрен и помещен в папку важных документов',
            category='success',
            link=url_for('documents.view_file', filename=pending_file.original_filename, category='important')
        )
        
        flash(f'Файл "{pending_file.original_filename}" успешно одобрен', 'success')
    except Exception as e:
        flash(f'Ошибка при одобрении файла: {str(e)}', 'error')
    
    return redirect(url_for('documents.pending_files'))

@documents_bp.route('/admin/reject_file/<int:file_id>', methods=['POST'])
@login_required
@role_required('admin')
def reject_file(file_id):
    from .models import db, PendingFile
    from .notifications import add_notification
    
    pending_file = PendingFile.query.get_or_404(file_id)
    
    if pending_file.status != 'pending':
        flash('Этот файл уже обработан', 'warning')
        return redirect(url_for('documents.pending_files'))
    
    try:
        # Обновляем статус файла
        pending_file.status = 'rejected'
        db.session.commit()
        
        # Отправляем уведомление пользователю
        add_notification(
            user_id=pending_file.user_id,
            message=f'Ваш запрос на добавление файла "{pending_file.original_filename}" в папку важных документов был отклонен',
            category='danger'
        )
        
        # Удаляем временный файл
        if os.path.exists(pending_file.temp_path):
            os.remove(pending_file.temp_path)
        
        flash(f'Файл "{pending_file.original_filename}" отклонен', 'info')
    except Exception as e:
        flash(f'Ошибка при отклонении файла: {str(e)}', 'error')
    
    return redirect(url_for('documents.pending_files'))
def secure_filename_with_cyrillic(filename):
    """
    Create a secure filename while preserving Cyrillic characters
    """
    # Convert spaces to underscores
    filename = filename.replace(' ', '_')
    
    # Keep only Cyrillic letters, Latin letters, numbers, and some special chars
    # This regex pattern allows Cyrillic (а-яА-ЯёЁ), Latin (a-zA-Z), digits, and some safe special chars
    cleaned_name = re.sub(r'[^а-яА-ЯёЁa-zA-Z0-9._-]', '', filename)
    
    # Remove any leading/trailing dots or spaces
    cleaned_name = cleaned_name.strip('._')
    
    # Ensure the filename isn't empty after cleaning
    if not cleaned_name:
        cleaned_name = 'unnamed_file'
        
    return cleaned_name

@documents_bp.route('/delete/<filename>', methods=['POST'])
@login_required
def delete_file(filename):
    try:
        category = request.form.get('category', 'personal')
        
        if category == 'personal':
            folder = get_user_upload_folder(current_user.id)
            trash_manager.move_to_trash(current_user.id, filename, os.path.join(folder, filename))
        elif category == 'common':
            folder = get_common_folder()
            trash_manager.move_to_trash('common', filename, os.path.join(folder, filename))
        elif category == 'important':
            folder = get_important_folder()
            trash_manager.move_to_trash('important', filename, os.path.join(folder, filename))
        
        flash(f"File moved to trash: {filename}", 'success')
        return redirect(url_for('documents.list_documents', category=category))
    except Exception as e:
        flash(f"Error moving file to trash: {str(e)}", 'error')
        return redirect(url_for('documents.list_documents', category='personal'))

@documents_bp.route('/view/<filename>')
@login_required
def view_file(filename):
    category = request.args.get('category', 'personal')
    
    if category == 'personal':
        folder = get_user_upload_folder(current_user.id)
    elif category == 'common':
        folder = get_common_folder()
    elif category == 'important':
        folder = get_important_folder()
    else:
        folder = get_user_upload_folder(current_user.id)
    
    return send_from_directory(folder, filename)

@documents_bp.route('/download/<filename>')
@login_required
def download_file(filename):
    category = request.args.get('category', 'personal')
    
    if category == 'personal':
        folder = get_user_upload_folder(current_user.id)
    elif category == 'common':
        folder = get_common_folder()
    elif category == 'important':
        folder = get_important_folder()
    else:
        folder = get_user_upload_folder(current_user.id)
    
    return send_from_directory(folder, filename, as_attachment=True)



def save_as_docx(content, filepath):
    """Сохранение текста как DOCX файла"""
    from docx import Document
    doc = Document()
    doc.add_paragraph(content)
    doc.save(filepath)

def save_as_pdf(content, filepath):
    """Сохранение текста как PDF файла"""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=content)
    pdf.output(filepath)
@documents_bp.route('/temp/preview/<filename>')
@login_required
@role_required('admin')
def preview_temp_file(filename):
    temp_folder = os.path.join(BASE_UPLOAD_FOLDER, 'pending')
    return send_from_directory(temp_folder, filename)

@documents_bp.route('/admin/user_documents/<int:user_id>')
@login_required
@role_required('admin')  # Only for administrators
def admin_view_user_documents(user_id):
    # Get the user to validate they exist
    from .models import User
    user = User.query.get_or_404(user_id)
    
    # Create the folder path for the user
    folder = get_user_upload_folder(user_id)
    
    documents = []
    
    try:
        for filename in os.listdir(folder):
            filepath = os.path.join(folder, filename)
            if os.path.isfile(filepath):
                file_stats = os.stat(filepath)
                documents.append({
                    'name': filename,
                    'size': f"{file_stats.st_size / 1024:.2f} KB",
                    'uploaded_at': datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M'),
                    'icon': get_file_icon(filename)
                })
    except Exception as e:
        flash(f'Error reading documents: {str(e)}', 'error')

    if not documents:
        documents = [{'name': "No documents available for this user."}]
    
    return render_template('documents/user_documents.html', documents=documents, user=user)
    